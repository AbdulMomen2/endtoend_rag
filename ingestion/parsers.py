"""
Document parsers with multimodal support.
Supports: PDF (text+images+tables), DOCX (text+tables), Excel/CSV, Images.
"""
import fitz
import docx
import base64
import logging
from abc import ABC, abstractmethod
from typing import List, Optional
from pathlib import Path
from langchain_core.documents import Document
from core.exceptions import UnsupportedFormatError, DocumentParsingError

logger = logging.getLogger(__name__)


def _describe_image(image_bytes: bytes, page_num: int) -> Optional[str]:
    """Use GPT-4o-mini vision to describe an image. Respects rate limits."""
    try:
        import openai, time
        client = openai.OpenAI()
        b64 = base64.b64encode(image_bytes).decode("utf-8")
        for attempt in range(3):
            try:
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": [
                        {"type": "text", "text": (
                            "Describe this image from a document in detail. "
                            "If it contains a chart, graph, or diagram, explain what it shows including all values, labels, and trends. "
                            "If it contains a table, extract all data as structured text. "
                            "Be precise and comprehensive."
                        )},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}}
                    ]}],
                    max_tokens=400
                )
                return response.choices[0].message.content.strip()
            except Exception as e:
                if "429" in str(e) and attempt < 2:
                    wait = 2 ** attempt  # 1s, 2s backoff
                    logger.warning(f"Rate limited on image p{page_num}, retrying in {wait}s")
                    time.sleep(wait)
                else:
                    raise
    except Exception as e:
        logger.warning(f"Image description failed for page {page_num}: {e}")
        return None


class BaseParser(ABC):
    @abstractmethod
    def parse(self, file_path: Path) -> List[Document]:
        pass


class PDFParser(BaseParser):
    def __init__(self, extract_images: bool = False):
        """
        extract_images: Set True to describe images via vision API.
        Disabled by default to avoid rate limits on large PDFs.
        Enable via EXTRACT_IMAGES=true env var or pass explicitly.
        """
        import os
        self.extract_images = extract_images or os.getenv("EXTRACT_IMAGES", "false").lower() == "true"

    def parse(self, file_path: Path) -> List[Document]:
        documents = []
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                page_number = page_num + 1
                text = page.get_text("text").strip()

                # Extract tables
                table_text = ""
                try:
                    tables = page.find_tables()
                    if tables and tables.tables:
                        for table in tables.tables:
                            df = table.to_pandas()
                            table_text += "\n[TABLE]:\n" + df.to_string(index=False) + "\n"
                except Exception:
                    pass

                full_text = text + table_text
                if full_text:
                    documents.append(Document(
                        page_content=full_text,
                        metadata={"source": str(file_path.name), "page": page_number, "has_tables": bool(table_text)}
                    ))

                # Extract and describe images
                if self.extract_images:
                    import time
                    for img_idx, img_info in enumerate(page.get_images(full=True)):
                        try:
                            xref = img_info[0]
                            base_image = doc.extract_image(xref)
                            img_bytes = base_image["image"]
                            if len(img_bytes) < 5000:
                                continue
                            time.sleep(0.5)  # Throttle to avoid rate limits
                            description = _describe_image(img_bytes, page_number)
                            if description:
                                documents.append(Document(
                                    page_content=f"[Image on Page {page_number}]: {description}",
                                    metadata={"source": str(file_path.name), "page": page_number, "type": "image"}
                                ))
                        except Exception as e:
                            logger.warning(f"Image extraction failed p{page_number}: {e}")

            doc.close()
            logger.info(f"Parsed PDF: {file_path.name} ({len(documents)} chunks)")
            return documents
        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {e}")
            raise DocumentParsingError(f"PDF parsing failed: {e}")


class DOCXParser(BaseParser):
    def parse(self, file_path: Path) -> List[Document]:
        try:
            doc = docx.Document(file_path)
            full_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            for i, table in enumerate(doc.tables):
                rows = [" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
                if rows:
                    full_text.append(f"\n[TABLE {i+1}]:\n" + "\n".join(rows))

            text_content = "\n".join(full_text)
            logger.info(f"Parsed DOCX: {file_path.name}")
            return [Document(page_content=text_content, metadata={"source": str(file_path.name), "type": "docx"})]
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise DocumentParsingError(f"DOCX parsing failed: {e}")


class ExcelParser(BaseParser):
    def parse(self, file_path: Path) -> List[Document]:
        try:
            import pandas as pd
            ext = file_path.suffix.lower()
            sheets = {"Sheet1": pd.read_csv(file_path)} if ext == ".csv" else {
                name: pd.ExcelFile(file_path).parse(name)
                for name in pd.ExcelFile(file_path).sheet_names
            }
            documents = []
            for sheet_name, df in sheets.items():
                if df.empty:
                    continue
                text = f"[Sheet: {sheet_name}]\n{df.to_string(index=False, max_rows=500)}"
                text += f"\n\n[Summary: {len(df)} rows x {len(df.columns)} columns. Columns: {', '.join(str(c) for c in df.columns)}]"
                documents.append(Document(
                    page_content=text,
                    metadata={"source": str(file_path.name), "sheet": sheet_name, "type": "spreadsheet"}
                ))
            logger.info(f"Parsed spreadsheet: {file_path.name} ({len(documents)} sheets)")
            return documents
        except ImportError:
            raise DocumentParsingError("Install pandas and openpyxl: pip install pandas openpyxl")
        except Exception as e:
            raise DocumentParsingError(f"Spreadsheet parsing failed: {e}")


class ImageParser(BaseParser):
    def parse(self, file_path: Path) -> List[Document]:
        try:
            img_bytes = file_path.read_bytes()
            description = _describe_image(img_bytes, 1)
            if not description:
                raise DocumentParsingError("Image description failed")
            return [Document(
                page_content=f"[Image: {file_path.name}]: {description}",
                metadata={"source": str(file_path.name), "type": "image", "page": 1}
            )]
        except Exception as e:
            raise DocumentParsingError(f"Image parsing failed: {e}")


class ParserFactory:
    SUPPORTED = {
        ".pdf": PDFParser,
        ".docx": DOCXParser,
        ".xlsx": ExcelParser,
        ".xls": ExcelParser,
        ".csv": ExcelParser,
        ".png": ImageParser,
        ".jpg": ImageParser,
        ".jpeg": ImageParser,
        ".webp": ImageParser,
    }

    @staticmethod
    def get_parser(file_path: Path) -> BaseParser:
        ext = file_path.suffix.lower()
        cls = ParserFactory.SUPPORTED.get(ext)
        if not cls:
            raise UnsupportedFormatError(
                f"Format '{ext}' not supported. Supported: {', '.join(ParserFactory.SUPPORTED.keys())}"
            )
        return cls()
